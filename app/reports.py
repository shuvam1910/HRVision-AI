import datetime
import io
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_pdf_report(df, high_risk_df, generated_by="HR Manager"):
    """Generates a premium executive PDF report summarizing HR Attrition analytics."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#4f46e5'), # Indigo
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#1f2937'),
        spaceBefore=15,
        spaceAfter=10
    )
    
    normal_style = ParagraphStyle(
        'DocNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#4b5563')
    )
    
    story = []
    
    # Document Header
    story.append(Paragraph("Smart HR Analytics & Attrition Report", title_style))
    
    # Metadata Block
    date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    metadata_text = f"<b>Generated On:</b> {date_str} | <b>Prepared By:</b> {generated_by} | <b>Status:</b> Executive Confidential"
    story.append(Paragraph(metadata_text, normal_style))
    story.append(Spacer(1, 15))
    story.append(Paragraph("This document contains automated predictions and analytics regarding organizational attrition risk. It identifies critical departments, summarizes employee statistics, lists high-risk targets, and supplies retention steps.", normal_style))
    story.append(Spacer(1, 20))
    
    # Section 1: Key Performance Metrics
    story.append(Paragraph("1. Organization Summary KPI Scorecard", h1_style))
    
    # Compute metrics
    total_emp = len(df)
    attrition_count = len(df[df["Attrition"] == "Yes"])
    attrition_rate = (attrition_count / total_emp) * 100
    avg_income = df["MonthlyIncome"].mean()
    avg_tenure = df["YearsAtCompany"].mean()
    
    kpi_data = [
        ["Metric KPI Indicator", "Value", "Status Reference"],
        ["Total Employees Analyzed", f"{total_emp:,}", "Baseline Roster"],
        ["Overall Attrition Rate", f"{attrition_rate:.1f}%", "Warning (>15% alert)" if attrition_rate > 15 else "Healthy Level"],
        ["Average Monthly Income", f"${avg_income:,.2f}", "Industry Competitive"],
        ["Average Tenure (Years)", f"{avg_tenure:.1f} Years", "Tenure Metric"]
    ]
    
    kpi_table = Table(kpi_data, colWidths=[200, 150, 150])
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4f46e5')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 8),
        ('TOPPADDING', (0,0), (-1,0), 8),
        ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#f9fafb')),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
        ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('TEXTCOLOR', (0,1), (0,-1), colors.HexColor('#1f2937')),
    ]))
    
    story.append(kpi_table)
    story.append(Spacer(1, 20))
    
    # Section 2: High Attrition Risk Roster
    story.append(Paragraph("2. Prioritized High-Risk Employees", h1_style))
    story.append(Paragraph("Roster of active employees predicted with high likelihood of leaving (Risk probability > 70%). Retention audits should prioritize this list.", normal_style))
    story.append(Spacer(1, 10))
    
    if len(high_risk_df) == 0:
        story.append(Paragraph("<i>No employees currently identified in the high-risk bracket.</i>", normal_style))
    else:
        roster_data = [["Emp ID", "Full Name", "Department", "Job Role", "Monthly Income", "Risk %"]]
        
        for _, row in high_risk_df.head(15).iterrows():
            roster_data.append([
                str(row["Employee ID"]),
                str(row["Full Name"]),
                str(row["Department"]),
                str(row["JobRole"]),
                f"${row['MonthlyIncome']:,}",
                f"{row['Attrition Risk (%)']}%"
            ])
            
        roster_table = Table(roster_data, colWidths=[50, 100, 100, 120, 80, 50])
        roster_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#ef4444')), # Red header
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('TOPPADDING', (0,0), (-1,0), 6),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(roster_table)
        
        if len(high_risk_df) > 15:
            story.append(Spacer(1, 5))
            story.append(Paragraph(f"<i>* Showing top 15 of {len(high_risk_df)} high-risk employees. Export CSV for the full report.</i>", normal_style))
            
    story.append(Spacer(1, 20))
    
    # Section 3: AI-Driven Action Guide
    story.append(Paragraph("3. Executive Action Guidelines", h1_style))
    story.append(Paragraph("Strategic recommendations to combat voluntary attrition based on machine learning drivers:", normal_style))
    story.append(Spacer(1, 8))
    
    recommendations_list = [
        "<b>Redistribute Overtime Commitments:</b> Overtime is statistically identified as the most severe driver of resignation. Assess workloads for employees working overtime regularly and adjust shift schedules or supply compensatory arrangements.",
        "<b>Market Rate Remuneration Audits:</b> Perform wage benchmarking for salaries falling below the $4,000/month threshold. Discrepancies between workload and pay in this category strongly drive attrition.",
        "<b>Incentive Plan Expansion:</b> Review the allocation of stock option programs. Zero-equity employee brackets present double the risk of departure compared to those with standard option incentives.",
        "<b>Management Relationship Workshops:</b> Guide department managers on communication and career pathway outlining. Work-life balance scores and manager-tenure dynamics require proactive coaching."
    ]
    
    for item in recommendations_list:
        bullet_text = f"• {item}"
        story.append(Paragraph(bullet_text, normal_style))
        story.append(Spacer(1, 6))
        
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_report(df, high_risk_df, generated_by="HR Manager"):
    """Generates a premium executive Microsoft Word (.docx) report summarizing HR Attrition analytics."""
    doc = docx.Document()
    
    # Styles set up
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Smart HR Analytics & Attrition Report")
    title_run.font.name = 'Helvetica'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo #4f46e5
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Metadata Block
    date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(f"Generated On: {date_str} | Prepared By: {generated_by} | Status: Executive Confidential")
    meta_run.font.name = 'Helvetica'
    meta_run.font.size = Pt(9.5)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(107, 114, 128) # Grey
    
    # Intro
    intro_p = doc.add_paragraph("This document contains automated predictions and analytics regarding organizational attrition risk. It identifies critical departments, summarizes employee statistics, lists high-risk targets, and supplies retention steps.")
    for run in intro_p.runs:
        run.font.name = 'Helvetica'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(55, 65, 81)
        
    doc.add_paragraph("") # Space
    
    # Section 1: KPI scorecard
    h1_p = doc.add_paragraph()
    h1_run = h1_p.add_run("1. Organization Summary KPI Scorecard")
    h1_run.font.name = 'Helvetica'
    h1_run.font.size = Pt(15)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(31, 41, 55)
    
    # Compute metrics
    total_emp = len(df)
    attrition_count = len(df[df["Attrition"] == "Yes"])
    attrition_rate = (attrition_count / total_emp) * 100
    avg_income = df["MonthlyIncome"].mean()
    avg_tenure = df["YearsAtCompany"].mean()
    
    # KPI Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Metric KPI Indicator", "Value", "Status Reference"]
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = title
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        
    kpi_rows = [
        ["Total Employees Analyzed", f"{total_emp:,}", "Baseline Roster"],
        ["Overall Attrition Rate", f"{attrition_rate:.1f}%", "Warning (>15% alert)" if attrition_rate > 15 else "Healthy Level"],
        ["Average Monthly Income", f"${avg_income:,.2f}", "Industry Competitive"],
        ["Average Tenure (Years)", f"{avg_tenure:.1f} Years", "Tenure Metric"]
    ]
    
    for r_idx, row in enumerate(kpi_rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9.5)
            
    doc.add_paragraph("") # Space
    
    # Section 2: High Attrition Risk Roster
    h2_p = doc.add_paragraph()
    h2_run = h2_p.add_run("2. Prioritized High-Risk Employees")
    h2_run.font.name = 'Helvetica'
    h2_run.font.size = Pt(15)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(31, 41, 55)
    
    p2 = doc.add_paragraph("Roster of active employees predicted with high likelihood of leaving (Risk probability > 70%). Retention audits should prioritize this list.")
    p2.runs[0].font.size = Pt(10.5)
    
    if len(high_risk_df) == 0:
        doc.add_paragraph("No employees currently identified in the high-risk bracket.")
    else:
        table_hr = doc.add_table(rows=1, cols=6)
        table_hr.style = 'Light Shading Accent 1'
        
        hr_headers = ["Emp ID", "Full Name", "Department", "Job Role", "Monthly Income", "Risk %"]
        hdr_cells = table_hr.rows[0].cells
        for i, header in enumerate(hr_headers):
            hdr_cells[i].text = header
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9.5)
            
        for _, row in high_risk_df.head(15).iterrows():
            row_cells = table_hr.add_row().cells
            row_cells[0].text = str(row["Employee ID"])
            row_cells[1].text = str(row["Full Name"])
            row_cells[2].text = str(row["Department"])
            row_cells[3].text = str(row["JobRole"])
            row_cells[4].text = f"${row['MonthlyIncome']:,}"
            row_cells[5].text = f"{row['Attrition Risk (%)']}%"
            for cell in row_cells:
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(9)
                
        if len(high_risk_df) > 15:
            p_note = doc.add_paragraph(f"* Showing top 15 of {len(high_risk_df)} high-risk employees. Export CSV for the full report.")
            p_note.runs[0].font.size = Pt(8.5)
            p_note.runs[0].font.italic = True
            
    doc.add_paragraph("") # Space
    
    # Section 3: AI Recommendations
    h3_p = doc.add_paragraph()
    h3_run = h3_p.add_run("3. Executive Action Guidelines")
    h3_run.font.name = 'Helvetica'
    h3_run.font.size = Pt(15)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(31, 41, 55)
    
    p3 = doc.add_paragraph("Strategic recommendations to combat voluntary attrition based on machine learning drivers:")
    p3.runs[0].font.size = Pt(10.5)
    
    recommendations = [
        ("Redistribute Overtime Commitments: ", "Overtime is statistically identified as the most severe driver of resignation. Assess workloads for employees working overtime regularly and adjust shift schedules or supply compensatory arrangements."),
        ("Market Rate Remuneration Audits: ", "Perform wage benchmarking for salaries falling below the $4,000/month threshold. Discrepancies between workload and pay in this category strongly drive attrition."),
        ("Incentive Plan Expansion: ", "Review the allocation of stock option programs. Zero-equity employee brackets present double the risk of departure compared to those with standard option incentives."),
        ("Management Relationship Workshops: ", "Guide department managers on communication and career pathway outlining. Work-life balance scores and manager-tenure dynamics require proactive coaching.")
    ]
    
    for title, desc in recommendations:
        bp = doc.add_paragraph(style='List Bullet')
        run_title = bp.add_run(title)
        run_title.font.bold = True
        run_title.font.size = Pt(10)
        
        run_desc = bp.add_run(desc)
        run_desc.font.size = Pt(10)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_single_employee_report(employee_data, prob, drivers, generated_by="HR Auditor"):
    """Generates a professional single-employee diagnostic PDF report."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=45,
        leftMargin=45,
        topMargin=45,
        bottomMargin=45
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'EmpTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'EmpH1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#1f2937'),
        spaceBefore=12,
        spaceAfter=8
    )
    
    normal_style = ParagraphStyle(
        'EmpNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#4b5563')
    )
    
    story = []
    
    story.append(Paragraph("Employee Attrition Risk Diagnostic", title_style))
    
    # Metadata
    date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    story.append(Paragraph(f"<b>Generated On:</b> {date_str} | <b>Auditor:</b> {generated_by} | <b>Status:</b> Strictly Confidential", normal_style))
    story.append(Spacer(1, 15))
    
    # Section 1: Employee Information
    story.append(Paragraph("1. Employee Information", h1_style))
    
    # Create Table of Info
    emp_details = [
        ["Age", str(employee_data.get("Age", "N/A")), "Gender", str(employee_data.get("Gender", "N/A"))],
        ["Department", str(employee_data.get("Department", "N/A")), "Job Role", str(employee_data.get("JobRole", "N/A"))],
        ["Monthly Income", f"${employee_data.get('MonthlyIncome', 0):,}", "Overtime Required", str(employee_data.get("OverTime", "N/A"))],
        ["Years at Company", f"{employee_data.get('YearsAtCompany', 0)} Yrs", "Years with Manager", f"{employee_data.get('YearsWithCurrManager', 0)} Yrs"]
    ]
    
    t_info = Table(emp_details, colWidths=[110, 140, 110, 140])
    t_info.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f3f4f6')),
        ('BACKGROUND', (2,0), (2,-1), colors.HexColor('#f3f4f6')),
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTNAME', (2,0), (2,-1), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
        ('PADDING', (0,0), (-1,-1), 6),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(t_info)
    story.append(Spacer(1, 15))
    
    # Section 2: Predictive Risk Metrics
    story.append(Paragraph("2. ML Attrition Prediction Results", h1_style))
    risk_label = "LOW RISK" if prob < 0.3 else ("MEDIUM RISK" if prob < 0.7 else "HIGH RISK")
    risk_color = "#10b981" if prob < 0.3 else ("#f59e0b" if prob < 0.7 else "#ef4444")
    
    story.append(Paragraph(f"Predicted Attrition Likelihood: <b><font color='{risk_color}'>{prob*100:.1f}%</font> ({risk_label})</b>", normal_style))
    story.append(Spacer(1, 10))
    
    # Section 3: Risk Driver Analysis
    story.append(Paragraph("3. Personalized Key Risk Drivers", h1_style))
    story.append(Paragraph("Factors driving the employee's predicted risk score up (+) or down (-):", normal_style))
    story.append(Spacer(1, 8))
    
    for impact, title, desc in drivers:
        color = "#ef4444" if "+" in impact else "#10b981"
        icon = "▲" if "+" in impact else "▼"
        bullet_html = f"• <b><font color='{color}'>{icon} {title} ({impact})</font></b>: {desc}"
        story.append(Paragraph(bullet_html, normal_style))
        story.append(Spacer(1, 6))
        
    story.append(Spacer(1, 15))
    
    # Section 4: AI Retention Action Plan
    story.append(Paragraph("4. Recommended Retention Action Plan", h1_style))
    recommendations = []
    for impact, title, desc in drivers:
        if "+" in impact:
            recommendations.append(f"Modify <b>{title}</b>: {desc}")
            
    if not recommendations:
        story.append(Paragraph("• No specific intervention is requested based on current metrics. Maintain standard checks.", normal_style))
    else:
        for rec in recommendations:
            story.append(Paragraph(f"• {rec}", normal_style))
            story.append(Spacer(1, 6))
            
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
